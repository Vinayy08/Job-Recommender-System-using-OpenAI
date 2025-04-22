import numpy as np
import pandas as pd
from django.core.paginator import Paginator
import csv , os
import re
import openai 
from functools import wraps
from django.db.models import Q
from django.db.models.signals import post_save
from django.dispatch import receiver
import spacy
import PyPDF2
from PyPDF2 import PdfReader
from docx import Document
from django.utils import timezone
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponse, HttpResponseForbidden ,JsonResponse ,Http404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .forms import (
    EmployerRegistrationForm,
    EmployeeRegistrationForm,
    LoginForm,
    EmployeeProfileForm,
    EmployerProfileForm,
    CSVUploadForm,
    JobForm,
    JobApplicationForm
)
from .models import User, UserProfile, Job, JobApplication
from .utils import (
    REPORTS_DIR,
    compare_resume_with_job,
    generate_detailed_compatibility_report,
    generate_employee_compatibility_report,
    recommend_top_jobs,   
    generate_clustered_bar_chart,
    generate_employee_clustered_chart,
    generate_recommendations,
    format_education,format_links,format_experience_projects
)
from .helpers import preprocess_text, extract_text_from_file
nlp = spacy.load("en_core_web_sm")
from main.decorators import employee_required
from django.conf import settings
import matplotlib.pyplot as plt
import seaborn as sns
import os
from django.urls import reverse
import matplotlib
import matplotlib.pyplot as plt
matplotlib.use('Agg')  # Use non-GUI backend for matplotlib

import logging
logger = logging.getLogger(__name__)


# Suppress matplotlib font matching logs
logging.getLogger('matplotlib.font_manager').disabled = True
logging.getLogger('matplotlib').setLevel(logging.ERROR)

# Suppress STREAM-related warnings
os.environ['PYTHONWARNINGS'] = 'ignore:STREAM b'


CACHE_DIR = os.path.join(settings.BASE_DIR, 'cache')
os.makedirs(CACHE_DIR, exist_ok=True)

# Helper decorator for employer access
def employer_required(view_func):
    @wraps(view_func)
    def _wrapped_view(request, *args, **kwargs):
        if hasattr(request.user, 'userprofile') and request.user.userprofile.role == 'employer':
            return view_func(request, *args, **kwargs)
        return HttpResponseForbidden("You are not authorized to access this page.")
    return _wrapped_view


# Helper decorator for employee access
def employee_required(view_func):
    @wraps(view_func)
    def _wrapped_view(request, *args, **kwargs):
        if hasattr(request.user, 'userprofile') and request.user.userprofile.role == 'employee':
            return view_func(request, *args, **kwargs)
        return HttpResponseForbidden("You are not authorized to access this page.")
    return _wrapped_view



@login_required
def home(request):
    try:
        profile = request.user.userprofile
        context = {
            'user': request.user,
            'full_name': profile.full_name,  # Make sure this is consistent with your model field names
            'email': request.user.email,
            'contact_number': profile.contact_number if profile.contact_number else "Not Provided",
        }
        if profile.role == 'employer':
            context.update({
                'employer_name': profile.full_name,
                'company_name': profile.company_name,
                'company_location': profile.company_location,
            })
            return render(request, 'main/employer_home.html', context)
        return render(request, 'main/employee_home.html', context)
    except UserProfile.DoesNotExist:
        # Display error on the current page instead of redirecting
        messages.error(request, "User profile not found. Please contact support.")
        # Provide a minimal context with just the user information
        context = {
            'user': request.user,
            'email': request.user.email,
        }
        # Render a simple error template or use a generic template
        return render(request, 'main/profile_error.html', context)
    

@login_required
def view_employee_profiles(request):
    # Get search query
    search_query = request.GET.get('search', '').strip()

    # Get employee profiles (excluding employers) with related user
    employee_profiles = UserProfile.objects.filter(role='employee').select_related('user').order_by('-experience_years')

    # Filter based on search query - expanded to more fields
    if search_query:
        employee_profiles = employee_profiles.filter(
            Q(full_name__icontains=search_query) |
            Q(skills__icontains=search_query) |
            Q(preferred_location__icontains=search_query) |
            Q(contact_email__icontains=search_query) |
            Q(contact_number__icontains=search_query) |
            Q(education__icontains=search_query) |
            Q(experience_years__icontains=search_query) |
            Q(experience_projects__icontains=search_query) |
            Q(links__icontains=search_query) |
            Q(testimonials__icontains=search_query) |
            Q(user__email__icontains=search_query)
        )

    # Format data before rendering
    formatted_profiles = []
    for profile in employee_profiles:
        profile.formatted_education = format_education(profile.education)
        profile.formatted_links = format_links(profile.links)
        profile.formatted_experience_projects = format_experience_projects(profile.experience_projects)
        formatted_profiles.append(profile)

    # Paginate the results
    paginator = Paginator(employee_profiles, 10)  # Show 10 profiles per page
    page_number = request.GET.get('page')
    profiles_page = paginator.get_page(page_number)

    return render(request, 'main/view_employee_profiles.html', {
        'employee_profiles': profiles_page,
        'search_query': search_query,
    })


@login_required
def get_certificate_details(request, certificate_id):
    """
    Fetch and return details of a specific certificate
    """
    try:
        certificate = EmployeeCertification.objects.get(id=certificate_id)
        
        # Prepare certificate details as a dictionary
        certificate_details = {
            'certificate_name': certificate.certificate_name,
            'issued_date': certificate.issued_date.strftime('%B %d, %Y'),
            'issuing_organization': certificate.issuing_organization,
            'description': certificate.description,
            'certificate_file': certificate.certificate_file.url if certificate.certificate_file else None
        }
        
        return JsonResponse(certificate_details)
    
    except EmployeeCertification.DoesNotExist:
        return JsonResponse({'error': 'Certificate not found'}, status=404)


@login_required
def employer_dashboard(request):
    search_query = request.GET.get('search', '').strip()  # Get the search term
    posted_jobs = Job.objects.filter(employer=request.user).order_by('-created_at')

    # Filter based on search query
    if search_query:
        posted_jobs = posted_jobs.filter(
            Q(company_name__icontains=search_query) |
            Q(job_description__icontains=search_query) |
            Q(role__icontains=search_query) |
            Q(industry_type__icontains=search_query) |
            Q(department__icontains=search_query) |
            Q(employment_type__icontains=search_query) |
            Q(role_category__icontains=search_query) |
            Q(education__icontains=search_query) |
            Q(skills__icontains=search_query) |
            Q(experience__icontains=search_query) |
            Q(location__icontains=search_query)
        )

    # Paginate the results
    paginator = Paginator(posted_jobs, 10)  # Show 10 jobs per page
    page_number = request.GET.get('page')
    jobs_page = paginator.get_page(page_number)

    return render(request, 'main/employer_dashboard.html', {
        'posted_jobs': jobs_page,
        'search_query': search_query,
    })


@login_required
@employer_required
def edit_job(request, job_id):
    job = get_object_or_404(Job, id=job_id, employer=request.user)
    if request.method == 'POST':
        form = JobForm(request.POST, instance=job)
        if form.is_valid():
            form.save()
            messages.success(request, "Job updated successfully!")
            return redirect('employer_dashboard')
    else:
        form = JobForm(instance=job)
    return render(request, 'main/edit_job.html', {'form': form, 'job': job})


@login_required
@employer_required
def delete_job(request, job_id):
    job = get_object_or_404(Job, id=job_id)
    
    # Store the company name before deletion
    company_name = job.company_name
    
    # Delete the job
    job.delete()
    
    # Add a success message with the company name
    messages.success(request, f'Job from "{company_name}" has been successfully deleted.')
    
    return redirect('employer_dashboard')


@login_required
@employee_required
def employee_dashboard(request):
    current_employee = request.user
    user_profile = current_employee.userprofile

    # Check for essential profile completeness (exclude experience_years and links)
    required_fields = [
        user_profile.full_name, 
        user_profile.resume, 
        user_profile.contact_number,
        user_profile.education,
        user_profile.skills,
        user_profile.experience_projects
    ]

    incomplete_profile = not all(required_fields)

    if incomplete_profile:
        messages.error(request, "Please complete your profile to access all features of the Employee Dashboard.")

    if not incomplete_profile:
        try:
            detailed_report = generate_employee_compatibility_report(current_employee)
            if not detailed_report:
                raise ValueError("No compatibility data found for this employee.")
            
            similarity_df = pd.DataFrame(
                {entry["Job"]: entry["Overall Compatibility"] / 100 for entry in detailed_report},
                index=["Overall Compatibility"]
            ).T

            clustered_chart_path = os.path.join(REPORTS_DIR, f'{user_profile.full_name}_compatibility_clustered_chart.png')
            if not os.path.exists(REPORTS_DIR):
                os.makedirs(REPORTS_DIR)
            generate_employee_clustered_chart(similarity_df["Overall Compatibility"], clustered_chart_path)

            resume_path = user_profile.resume.url if user_profile.resume and user_profile.resume.name else None

            return render(request, 'main/employee_dashboard.html', {
                'clustered_chart_path': f'/static/reports/{os.path.basename(clustered_chart_path)}',
                'employee_name': user_profile.full_name,
                'resume_path': resume_path,
                'detailed_report': detailed_report,
                'incomplete_profile': incomplete_profile
            })

        except Exception as e:
            logger.error(f"Error in employee_dashboard: {e}")
            return render(request, 'main/error.html', {"error_message": str(e)})

    return render(request, 'main/employee_dashboard.html', {
        'employee_name': user_profile.full_name,
        'incomplete_profile': incomplete_profile
    })


from django.db import IntegrityError

from django.views.decorators.csrf import csrf_exempt, csrf_protect
from django.utils.decorators import method_decorator


# Employer Registration View
@csrf_protect
def employer_register(request):
    registration_success = False
    if request.method == 'POST':
        form = EmployerRegistrationForm(request.POST)
        if form.is_valid():
            try:
                user = form.save()
                registration_success = True
                # No message needed as we're using the template variable
                # No immediate redirect - let JavaScript handle it
            except IntegrityError:
                messages.error(request, "A user with this email or contact number already exists.")
            except Exception as e:
                messages.error(request, f"An unexpected error occurred during registration: {e}")
        else:
            messages.error(request, "Please correct the errors below.")
    else:
        form = EmployerRegistrationForm()
    return render(request, 'main/employer_register.html', {
        'form': form,
        'registration_success': registration_success
    })

# Employee Registration View
@csrf_protect
def employee_register(request):
    registration_success = False
    if request.method == 'POST':
        form = EmployeeRegistrationForm(request.POST)
        if form.is_valid():
            try:
                user = form.save()
                registration_success = True
                # No message needed as we're using the template variable
                # No immediate redirect - let JavaScript handle it
            except Exception as e:
                messages.error(request, f"Unexpected error during registration: {e}")
        else:
            messages.error(request, "Please correct the errors below.")
    else:
        form = EmployeeRegistrationForm()
    return render(request, 'main/employee_register.html', {
        'form': form,
        'registration_success': registration_success
    })


@login_required
@employer_required
def update_employer_profile(request):
    user_profile = get_object_or_404(UserProfile, user=request.user)

    if request.method == 'POST':
        form = EmployerProfileForm(request.POST, instance=user_profile, user_instance=request.user)
        if form.is_valid():
            form.save()
            return redirect('employer_dashboard')
    else:
        form = EmployerProfileForm(instance=user_profile, user_instance=request.user)

    return render(request, 'main/update_employer_profile.html', {
        'form': form,
        'user_profile': user_profile,
    })

from django.contrib.auth.models import User
from django.contrib.auth import get_user_model

User = get_user_model()


# Update the update_employee_profile view to ensure extracted data is properly formatted

@login_required
@employee_required
def update_employee_profile(request):
    user_profile = get_object_or_404(UserProfile, user=request.user)
    form_data = None

    # Format any JSON strings that might be stored in the database before rendering
    if isinstance(user_profile.education, str) and (user_profile.education.startswith('[') and user_profile.education.endswith(']')):
        user_profile.education = format_education(user_profile.education)
        
    # Handle links formatting - ensure empty links are truly empty
    formatted_links = format_links(user_profile.links)
    if not formatted_links:
        user_profile.links = ""
    else:
        user_profile.links = formatted_links
        
    if isinstance(user_profile.experience_projects, str) and (user_profile.experience_projects.startswith('[') and user_profile.experience_projects.endswith(']')):
        user_profile.experience_projects = format_experience_projects(user_profile.experience_projects)
        
    if request.method == 'POST':
        # Store the POST data in case we need to re-render the form with errors
        form_data = request.POST.copy()
        
        # Existing POST handling code remains the same
        name = request.POST.get('name', '').strip()
        email = request.POST.get('email', '').strip()
        contact_number = request.POST.get('contact_number', '').strip()
        links = request.POST.get('links', '').strip()
        skills = request.POST.get('skills', '').strip()
        education = request.POST.get('education', '').strip()
        experience_projects = request.POST.get('experience_projects', '').strip()
        preferred_location = request.POST.get('preferred_location', '').strip()
        testimonials = request.POST.get('testimonials', '').strip()
        resume = request.FILES.get('resume', None)
        old_password = request.POST.get('old_password', '').strip()
        new_password = request.POST.get('new_password', '').strip()
        
      
        # Safe conversion of expected_salary
        try:
            expected_salary = request.POST.get('expected_salary')
            if expected_salary:
                expected_salary = int(expected_salary)
                user_profile.expected_salary = expected_salary
        except ValueError:
            messages.error(request, "Invalid expected salary entered. Please enter a valid number.")
            # Return the form with the original data instead of redirecting
            return render(request, 'main/update_profile.html', {
                'user_profile': user_profile,
                'form_data': form_data
            })

        try:
            if name:
                user_profile.full_name = name
            if email and email != request.user.email:
                if User.objects.filter(email=email).exists():
                    messages.error(request, "Email already in use.")
                    # Return the form with the original data instead of redirecting
                    return render(request, 'main/update_profile.html', {
                        'user_profile': user_profile,
                        'form_data': form_data
                    })
                request.user.email = email

            if contact_number and contact_number != user_profile.contact_number:
                if UserProfile.objects.filter(contact_number=contact_number).exists():
                    messages.error(request, "Contact number already in use.")
                    # Return the form with the original data instead of redirecting
                    return render(request, 'main/update_profile.html', {
                        'user_profile': user_profile,
                        'form_data': form_data
                    })
                user_profile.contact_number = contact_number

            user_profile.links = links
            user_profile.skills = skills
            user_profile.education = education
            user_profile.experience_projects = experience_projects
            user_profile.preferred_location = preferred_location
            user_profile.testimonials = testimonials
            
            # Process resume if uploaded
            if resume:
                try:
                    # Parse the resume
                    resume_text = parse_resume(resume)
                    # Extract data using OpenAI
                    extracted_data = extract_data_from_resume(resume_text)
                    
                    # Store the resume file
                    user_profile.resume = resume
                    
                    # Update fields with extracted data, but only if the extracted data is not empty
                    if extracted_data.get('skills'):
                        user_profile.skills = extracted_data.get('skills')
                    if extracted_data.get('education'):
                        user_profile.education = extracted_data.get('education')
                    if extracted_data.get('experience_projects'):
                        user_profile.experience_projects = extracted_data.get('experience_projects')
                    if extracted_data.get('experience_years'):
                        user_profile.experience_years = extracted_data.get('experience_years')
                    if extracted_data.get('contact_number'):
                        user_profile.contact_number = extracted_data.get('contact_number')
                    
                    # Specifically handle links
                    if extracted_data.get('links') is not None:
                        formatted_links = format_links(extracted_data.get('links'))
                        if formatted_links:
                            user_profile.links = formatted_links
                        else:
                            user_profile.links = ""  # Set to empty string if no links or empty list
                    
                    if extracted_data.get('preferred_location'):
                        user_profile.preferred_location = extracted_data.get('preferred_location')
                    if extracted_data.get('testimonials'):
                        user_profile.testimonials = extracted_data.get('testimonials')
                    if extracted_data.get('expected_salary'):
                        user_profile.expected_salary = extracted_data.get('expected_salary')
                    
                    # If name is extracted and user didn't provide one manually, use the extracted name
                    if not name and extracted_data.get('name'):
                        user_profile.full_name = extracted_data.get('name')
                    
                    # If email is extracted and user didn't provide one manually
                    if not email and extracted_data.get('email') and extracted_data.get('email') != request.user.email:
                        # Check if email is already in use
                        if not User.objects.filter(email=extracted_data.get('email')).exists():
                            request.user.email = extracted_data.get('email')
                    
                    messages.success(request, "Resume uploaded and data extracted successfully.")
                except Exception as e:
                    messages.warning(request, f"Resume uploaded but data extraction had issues: {str(e)}")

            if old_password and new_password:
                if request.user.check_password(old_password):
                    request.user.set_password(new_password)
                    request.user.save()
                    messages.success(request, "Password changed successfully. Please log in again.")
                    logout(request)
                    return redirect('login')
                else:
                    messages.error(request, "Incorrect old password.")
                    # Return the form with the original data instead of redirecting
                    return render(request, 'main/update_profile.html', {
                        'user_profile': user_profile,
                        'form_data': form_data
                    })

            # Before saving, make sure the data is properly formatted
            if isinstance(user_profile.education, str) and (user_profile.education.startswith('[') and user_profile.education.endswith(']')):
                user_profile.education = format_education(user_profile.education)
                
            # Final check for links to ensure empty links are truly empty
            formatted_links = format_links(user_profile.links)
            if not formatted_links:
                user_profile.links = ""
            else:
                user_profile.links = formatted_links
                
            if isinstance(user_profile.experience_projects, str) and (user_profile.experience_projects.startswith('[') and user_profile.experience_projects.endswith(']')):
                user_profile.experience_projects = format_experience_projects(user_profile.experience_projects)

            user_profile.save()
            request.user.save()
            messages.success(request, "Profile updated successfully.")
            # Redirect to the same page after successful update
            return redirect('update_profile')

        except Exception as e:
            messages.error(request, f"An error occurred: {str(e)}")
            # Return the form with the original data instead of redirecting
            return render(request, 'main/update_profile.html', {
                'user_profile': user_profile,
                'form_data': form_data
            })

    return render(request, 'main/update_profile.html', {'user_profile': user_profile})


import json
import re
import traceback
from openai import OpenAI
        
def parse_resume(file):
    """
    Enhanced resume text extraction with better error handling
    """
    import PyPDF2
    from docx import Document
    import io
    
    content = ""
    try:
        # Create a copy of the file in memory to avoid file pointer issues
        file_copy = io.BytesIO(file.read())
        file.seek(0)  # Reset the file pointer
        
        if file.name.lower().endswith(".pdf"):
            reader = PyPDF2.PdfReader(file_copy)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n"
                    
        elif file.name.lower().endswith(".docx"):
            doc = Document(file_copy)
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
        else:
            raise ValueError(f"Unsupported file format: {file.name}")
            
        # Clean up the extracted text
        content = content.replace("Internal", "")  # Remove watermarks
        content = ' '.join(content.split())  # Normalize whitespace
        
    except Exception as e:
        print(f"Error reading file {file.name}: {str(e)}")
    
    print(f"Extracted {len(content)} characters from resume")
    return content

def extract_data_from_resume(resume_text):
    """
    Extract structured data from resume text using OpenAI's API with gpt-3.5-turbo
    to minimize token usage and costs.
    """
    import re
    import json
    from urllib.parse import urlparse

    try:
        # Get API key from settings
        openai_api_key = getattr(settings, 'OPENAI_API_KEY', None)
        if not openai_api_key:
            print("OpenAI API key is not configured.")
            return {}
        
        # Set up OpenAI client
        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        
        # Create a concise prompt to reduce token usage - Fixed the problematic formatting
        prompt = """
        Extract the following information from this professional resume as JSON. 
        Be thorough and look for information across the entire document. 
        Fields to extract:
        - name: The candidate's full name (e.g., "John Doe")
        - email: Email address (look for @ symbol, e.g., "example@email.com")
        - contact_number: Phone number with country code if available (e.g., "+91-1234567890")
        - skills: ALL technical skills mentioned (databases, languages, tools, platforms) as comma separated list
        - links: Any LinkedIn, GitHub or other professional URLs as an array of objects with platform and url fields
        - education: All education details including degrees, institutions, years
        - experience_projects: Summary of work experience with companies and roles
        - experience_years: Total years of professional experience as a number (e.g., 10)
        - preferred_location: Current or preferred location if mentioned
        - testimonials: Any testimonials or recommendations if present
        - expected_salary: Expected salary if mentioned (number only)
        
        For the links field, return an array of objects with platform and url fields like:
        [{"platform": "LinkedIn", "url": "https://linkedin.com/in/username"}]
        IMPORTANT: If no links are found, return an empty array for 'links' field. DO NOT include objects with empty URLs.

        For the given resume, do a thorough search and extract all information.
        
        Resume:
        """ + resume_text
        
        # Increase token limit for more complete extraction
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert resume parser. Extract all relevant information completely and accurately."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,  # Lower temperature for consistency
            max_tokens=1000  # Increased token limit
        )
        
        # Parse the JSON response
        print("OpenAI Response:", response.choices[0].message.content)
        text = (response.choices[0].message.content)
        clean_text = re.sub(r"```(?:json|python)?\n?", "", text).strip()
        clean_text = clean_text.replace("```", "").strip()
        # print("result2=======",clean_text)
        # print("type=============",type(response.choices[0].message.content))
        result = json.loads(clean_text)
        
        # Ensure all expected fields are present
        expected_fields = ["name", "email", "contact_number", "skills", "links", "education", 
                          "experience_projects", "experience_years", "preferred_location", 
                          "testimonials", "expected_salary"]
        
        for field in expected_fields:
            if field not in result:
                result[field] = "" if field != "experience_years" and field != "expected_salary" and field != "links" else 0 if field != "links" else []
                
        # Convert numeric fields to proper types
        try:
            if result["experience_years"] and not isinstance(result["experience_years"], int):
                result["experience_years"] = int(float(str(result["experience_years"]).replace(',', '')))
        except (ValueError, TypeError):
            result["experience_years"] = 0
            
        try:
            if result["expected_salary"] and not isinstance(result["expected_salary"], int):
                # Handle potential currency symbols or 'k' suffix
                salary_str = str(result["expected_salary"]).replace(',', '')
                # Remove non-numeric characters except decimal point
                salary_str = ''.join(c for c in salary_str if c.isdigit() or c == '.')
                if salary_str:
                    # Check for 'k' notation in the original string
                    if 'k' in str(result["expected_salary"]).lower():
                        result["expected_salary"] = int(float(salary_str) * 1000)
                    else:
                        result["expected_salary"] = int(float(salary_str))
        except (ValueError, TypeError):
            result["expected_salary"] = 0

        # Add a fallback URL scanner if links are empty or malformed
        # First ensure links is a list
        if not isinstance(result["links"], list):
            result["links"] = []
            
        # Only scan for links if the list is empty
        if len(result["links"]) == 0:
            # Look for common professional URLs in the resume text
            linkedin_pattern = r'linkedin\.com/\S+'
            github_pattern = r'github\.com/\S+'
            website_pattern = r'https?://(?!linkedin\.com|github\.com)(?:[-\w.]|(?:%[\da-fA-F]{2}))+[/\w\.-]*'
            
            found_links = []
            
            # Find LinkedIn profiles
            linkedin_matches = re.findall(linkedin_pattern, resume_text, re.IGNORECASE)
            for match in linkedin_matches:
                url = match if match.startswith('http') else f"https://{match}"
                found_links.append({"platform": "LinkedIn", "url": url})
            
            # Find GitHub profiles
            github_matches = re.findall(github_pattern, resume_text, re.IGNORECASE)
            for match in github_matches:
                url = match if match.startswith('http') else f"https://{match}"
                found_links.append({"platform": "GitHub", "url": url})
            
            # Find other websites
            website_matches = re.findall(website_pattern, resume_text, re.IGNORECASE)
            for match in website_matches:
                parsed_url = urlparse(match)
                domain = parsed_url.netloc.replace('www.', '')
                platform = domain.split('.')[0].capitalize()
                found_links.append({"platform": platform, "url": match})
            
            # Only assign found_links if we actually found some links
            if found_links:
                result["links"] = found_links
        
        # Clean up the links list to remove any objects with empty URLs
        if isinstance(result["links"], list):
            result["links"] = [link for link in result["links"] if link.get("url") and link.get("url").strip()]
            
            # Check if any platform has an empty URL, and if so, remove it from the links list
            for link in result["links"]:
                if not link.get("url") or not link.get("url").strip():
                    result["links"].remove(link)
        
        print("Extracted data:", result)
        return result
        
    except Exception as e:
        print(f"Error in OpenAI resume parsing: {str(e)}")
        import traceback
        traceback.print_exc()
        # Return empty values as fallback
        return {
            "name": "",
            "email": "",
            "contact_number": "",
            "skills": "",
            "links": [],  # Empty array, not objects with empty URLs
            "education": "",
            "experience_projects": "",
            "experience_years": 0,
            "preferred_location": "",
            "testimonials": "",
            "expected_salary": 0
        }
        
def post_process_result(result):
    """Format and clean up extracted data for better display"""
    # Handle expected_salary
    if isinstance(result["expected_salary"], str):
        try:
            import re
            digits = re.sub(r'[^\d.]', '', result["expected_salary"])
            if digits:
                if 'k' in result["expected_salary"].lower():
                    result["expected_salary"] = int(float(digits) * 1000)
                else:
                    result["expected_salary"] = int(float(digits))
            else:
                result["expected_salary"] = 0
        except (ValueError, TypeError):
            result["expected_salary"] = 0
    elif result["expected_salary"] is None:
        result["expected_salary"] = 0
    
    # Convert None values to appropriate defaults
    for key in result:
        if result[key] is None:
            if key in ["experience_years", "expected_salary"]:
                result[key] = 0
            elif isinstance(result[key], list):
                result[key] = []  # Empty list
            else:
                result[key] = ""
    
    # Ensure experience_years is an integer
    if isinstance(result["experience_years"], str):
        try:
            import re
            years_text = re.sub(r'[^\d.]', '', result["experience_years"])
            if years_text:
                result["experience_years"] = int(float(years_text))
            else:
                result["experience_years"] = 0
        except (ValueError, TypeError):
            result["experience_years"] = 0
    elif isinstance(result["experience_years"], (int, float)):
        result["experience_years"] = int(result["experience_years"])
    else:
        result["experience_years"] = 0
    
    # Format data fields without HTML tags
    result["links"] = format_links(result["links"])
    result["education"] = format_education(result["education"])
    result["experience_projects"] = format_experience_projects(result["experience_projects"])
    
    return result

@login_required
@employer_required
def upload_csv(request):
    if request.method == 'POST':
        form = CSVUploadForm(request.POST, request.FILES)
        if form.is_valid():
            csv_file = request.FILES['csv_file']
            decoded_file = csv_file.read().decode('utf-8').splitlines()
            reader = csv.DictReader(decoded_file)

            added_count = 0
            skipped_count = 0
            existing_jobs = set(
                Job.objects.filter(employer=request.user).values_list('company_name', 'role', 'location', named=True)
            )

            for row in reader:
                job_key = (row['Company Name'], row['Role'], row['Location'])
                if job_key not in existing_jobs:
                    Job.objects.create(
                        employer=request.user,
                        company_name=row['Company Name'],
                        job_description=row['Job Description'],
                        role=row['Role'],
                        industry_type=row['Industry Type'],
                        department=row['Department'],
                        employment_type=row['Employment Type'],
                        role_category=row['Role Category'],
                        education=row['Education'],
                        skills=row['Skills'],
                        experience=row['Experience'],
                        location=row['Location'],
                    )
                    added_count += 1
                else:
                    skipped_count += 1

            messages.success(
                request,
                f"{added_count} jobs added successfully. {skipped_count} duplicate jobs were skipped."
            )
            return redirect('employer_dashboard')
        else:
            for error in form.errors.values():
                messages.error(request, error.as_text())
    else:
        form = CSVUploadForm()

    return render(request, 'main/upload_csv.html', {'form': form})


def login_user(request):
    if request.method == "POST":
        form = LoginForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data.get("email")
            password = form.cleaned_data.get("password")
            user = authenticate(request, username=email, password=password)
            if user is not None:
                login(request, user)
                return redirect("home")
            else:
                form.add_error(None, "Invalid email or password")
    else:
        form = LoginForm()
    return render(request, "main/login.html", {"form": form})

def logout_user(request):
    logout(request)
    return redirect('login')



@login_required
@employer_required
def add_job(request):
    if request.method == 'POST':
        form = JobForm(request.POST)
        if form.is_valid():
            company_name = form.cleaned_data.get('company_name')
            role = form.cleaned_data.get('role')
            location = form.cleaned_data.get('location')

            if Job.objects.filter(
                employer=request.user,
                company_name=company_name,
                role=role,
                location=location,
            ).exists():
                messages.error(request, "This job already exists.")
            else:
                job = form.save(commit=False)
                job.employer = request.user
                job.save()
                messages.success(request, "Job posted successfully!")
                return redirect('employer_dashboard')
        else:
            for error in form.errors.values():
                messages.error(request, error.as_text())
    else:
        form = JobForm()

    return render(request, 'main/add_job.html', {'form': form})


@login_required
@employee_required
def view_all_jobs(request):
    """
    View all jobs with search functionality and applied job tracking.
    """
    search_query = request.GET.get('search', '').strip()
    jobs = Job.objects.all()

    # Search across all relevant fields
    if search_query:
        jobs = jobs.filter(
            Q(company_name__icontains=search_query) |
            Q(job_description__icontains=search_query) |
            Q(role__icontains=search_query) |
            Q(industry_type__icontains=search_query) |
            Q(department__icontains=search_query) |
            Q(employment_type__icontains=search_query) |
            Q(role_category__icontains=search_query) |
            Q(education__icontains=search_query) |
            Q(skills__icontains=search_query) |
            Q(experience__icontains=search_query) |
            Q(location__icontains=search_query)
        )

    # Get a set of job IDs the employee has applied for
    applied_job_ids = set(JobApplication.objects.filter(user=request.user).values_list('job_id', flat=True))

    return render(request, 'main/view_jobs.html', {'jobs': jobs, 'applied_job_ids': applied_job_ids})


@login_required
@employee_required
def view_recommendations(request):
    top_10_scores = request.session.get('top_10_compatibility_scores', [])
    if not top_10_scores:
        messages.warning(request, "Please calculate compatibility scores to view Top 10 Job Recommendations.")
        return render(request, 'main/view_recommendations.html', {'recommended_jobs': []})

    top_10_jobs = []
    for slug, data in top_10_scores:
        company_name = data['company_name']
        compatibility_score = data['score']
        job_details = Job.objects.filter(company_name__iexact=company_name).first()
        if job_details:
            top_10_jobs.append({
                'id': job_details.id,
                'company_name': job_details.company_name,
                'role': job_details.role,
                'compatibility_score': compatibility_score
            })

    applied_job_ids = JobApplication.objects.filter(user=request.user).values_list('job_id', flat=True)
    return render(request, 'main/view_recommendations.html', {
        'recommended_jobs': top_10_jobs,
        'applied_job_ids': applied_job_ids,
    })


@login_required
@employee_required
def job_detail(request, job_id):
    job = get_object_or_404(Job, id=job_id)
    return render(request, 'main/job_detail.html', {'job': job})


from django.http import JsonResponse

@login_required
@employee_required
def apply_for_job(request, job_id):
    job = get_object_or_404(Job, id=job_id)
    user_profile = request.user.userprofile

    # Check if the user has already applied
    if JobApplication.objects.filter(user=request.user, job=job).exists():
        return JsonResponse({'message': 'Already applied'}, status=400)

    # Process the uploaded resume
    resume = request.FILES.get('resume') or user_profile.resume

    # Create the job application
    JobApplication.objects.create(
        user=request.user,
        job=job,
        user_profile=user_profile,
        resume=resume
    )
    return JsonResponse({'message': 'Application successful'}, status=200)



from django.db.models import Prefetch
from main.models import EmployeeCertification

@login_required
@employer_required
def view_all_applications(request):
    """
    View all applications across all jobs posted by the logged-in employer.
    Optimized with select_related and prefetch_related for user, profile, and certifications.
    """
    applications = JobApplication.objects.filter(job__employer=request.user).select_related(
        'job', 'user__userprofile'
    ).prefetch_related(
        Prefetch('user__certifications', queryset=EmployeeCertification.objects.all())
    )
    return render(request, 'main/view_applications.html', {'applications': applications})


@login_required
@employer_required
def view_applications(request, job_id=None):
    """
    View applications for a specific job (if job_id is provided) or all jobs for the employer.
    """
    try:
        if job_id:
            applications = JobApplication.objects.filter(
                job__id=job_id,
                job__employer=request.user
            ).select_related(
                'job', 'user__userprofile'
            ).prefetch_related(
                Prefetch('user__certifications', queryset=EmployeeCertification.objects.all())
            )
        else:
            applications = JobApplication.objects.filter(
                job__employer=request.user
            ).select_related(
                'job', 'user__userprofile'
            ).prefetch_related(
                Prefetch('user__certifications', queryset=EmployeeCertification.objects.all())
            )

        return render(request, 'main/view_applications.html', {'applications': applications})
    except Exception as e:
        return render(request, 'main/error.html', {"error_message": str(e)})



@login_required
@employee_required
def view_employee_applications(request):
    """
    View applications submitted by the logged-in employee.
    """
    # Fetch applications for the logged-in employee
    applications = JobApplication.objects.filter(user=request.user).select_related('job')

    return render(request, 'main/view_employee_applications.html', {'applications': applications})


def download_resume(request, application_id):
    application = get_object_or_404(JobApplication, id=application_id)
    if not application.resume or not application.resume.path:
        raise Http404("Resume file not found.")
    return redirect(application.resume.url)
    

@login_required
@employer_required
def delete_application_employer(request, application_id):
    """
    Delete a job application by an employer.
    """
    try:
        application = get_object_or_404(JobApplication, id=application_id)

        # Debugging output to include full name
        print(f"Logged-in user: {request.user.userprofile.full_name} (ID: {request.user.id})")
        print(f"Job's employer: {application.job.employer.userprofile.full_name} (ID: {application.job.employer.id})")

        # Compare by ID to avoid object mismatch
        if application.job.employer.id != request.user.id:
            print(f"Authorization failed for user {request.user.userprofile.full_name} on job owned by {application.job.employer.userprofile.full_name}")
            return render(request, 'main/error.html', {"error_message": "You are not authorized to access this page."})

        # If authorized, delete the application
        application.delete()
        messages.success(request, "Application deleted successfully.")
        return redirect('view_applications', job_id=application.job.id)

    except JobApplication.DoesNotExist:
        return render(request, 'main/error.html', {"error_message": "Application does not exist."})


@login_required
@employee_required
def delete_application_employee(request, application_id):
    """
    Delete a job application submitted by the employee.
    """
    try:
        # Check if the application belongs to the logged-in employee
        application = get_object_or_404(JobApplication, id=application_id, user=request.user)

        # Delete the application
        application.delete()
        messages.success(request, "Your application has been successfully deleted.")
        return redirect('view_employee_applications')

    except JobApplication.DoesNotExist:
        return render(request, 'main/error.html', {"error_message": "Application not found or you do not have permission to delete it."})


@login_required
def view_resume(request):
    user_profile = request.user.userprofile  # Assuming a one-to-one relation between User and UserProfile
    resume_path = user_profile.resume.path if user_profile.resume else None

    if resume_path and os.path.exists(resume_path):
        with open(resume_path, 'rb') as resume_file:
            response = HttpResponse(resume_file.read(), content_type="application/pdf")
            response['Content-Disposition'] = f'inline; filename="{os.path.basename(resume_path)}"'
            return response
    return render(request, 'main/error.html', {"error_message": "Resume not found."})


@login_required
def view_employer_compatibility(request, company, user_id):
    try:
        job = get_object_or_404(Job, company_name=company)
        user_profile = get_object_or_404(UserProfile, user_id=user_id, role='employee')

        # Determine the email to display
        contact_email = (
            user_profile.contact_email or  # First, try UserProfile's contact_email
            user_profile.user.email  # Fallback to the associated User's email
        )

        user_profile = UserProfile.objects.select_related('user').prefetch_related('user__certifications').get(user_id=user_id, role='employee')


        # Refresh job and user profile data from the database to ensure they are up-to-date
        job.refresh_from_db()
        user_profile.refresh_from_db()

        job_details = {
            "education": job.education or "Not Specified",
            "experience": job.experience or "0 - 0 years",
            "skills": [skill.strip() for skill in (job.skills or "").split(",") if skill.strip()] or ["Not Specified"]
        }

        resume_details = {
            "education": format_education(user_profile.education) or "Not Specified",
            "experience_years": user_profile.experience_years or 0,
            "skills": [skill.strip() for skill in (user_profile.skills or "").split(",") if skill.strip()] or ["Not Specified"]
        }

        try:
            compatibility_matrix, scores = compare_resume_with_job(resume_details, job_details)
            overall_compatibility = scores  # Pass the entire scores dictionary
        except Exception as e:
            logger.error(f"Failed to process compatibility data: {e}")
            return render(request, 'main/error.html', {"error_message": "Failed to process compatibility data."})

        recommendations = generate_recommendations(compatibility_matrix, job_details)

        return render(request, 'main/employer_side_compatibility_display.html', {
            "employee_full_name": user_profile.full_name,
            "company_name": company,
            "compatibility_matrix": compatibility_matrix,
            "overall_compatibility": overall_compatibility,
            "recommendations": recommendations,
            "user_profile": user_profile,
            "contact_email": contact_email,
            "contact_number": user_profile.contact_number or "N/A"
        })
    except Exception as e:
        logger.error(f"Error in view_employer_compatibility with {company} and {user_id}: {e}", exc_info=True)
        return render(request, 'main/error.html', {"error_message": f"An unexpected error occurred: {str(e)}"})


@login_required
def view_compatibility_report(request, job_id):
    """
    View Compatibility Report for a specific job posting.
    """
    try:
        employer = request.user
        job = Job.objects.get(id=job_id, employer=employer)

        # Generate detailed compatibility report
        detailed_report = generate_detailed_compatibility_report(employer=employer, job_id=job_id)

        # Prepare data for the similarity matrix table and chart
        similarity_data = {
            entry["Candidate"]: {"Company": entry["Job"], "Overall Compatibility": entry["Overall Compatibility"] / 100}
            for entry in detailed_report
        }

        # Generate a DataFrame for visualization
        similarity_df = pd.DataFrame(
            {
                candidate: data["Overall Compatibility"] for candidate, data in similarity_data.items()
            },
            index=["Overall Compatibility"]
        ).T

        # Create the bar chart visualization
        bar_chart_path = os.path.join(
            settings.STATICFILES_DIRS[0], f'reports/compatibility_report_job_{job_id}.png'
        )
        generate_clustered_bar_chart(similarity_df, bar_chart_path)

        # Convert similarity data into a format suitable for the table
        similarity_matrix_table = """
        <thead>
            <tr>
                <th>Candidate</th>
                <th>Company</th>
                <th>Overall Compatibility</th>
            </tr>
        </thead>
        <tbody>
        """
        for candidate, data in similarity_data.items():
            similarity_matrix_table += f"<tr><td>{candidate}</td><td>{data['Company']}</td><td>{data['Overall Compatibility']:.2f}</td></tr>"
        similarity_matrix_table += "</tbody>"

        return render(request, 'main/compatibility_report.html', {
            "job": job,
            "bar_chart_path": f"/static/reports/compatibility_report_job_{job_id}.png",
            "similarity_matrix": similarity_matrix_table,
        })

    except Job.DoesNotExist:
        return render(request, 'main/error.html', {"error_message": "Job not found or you do not have permission to view it."})
    except Exception as e:
        logger.error(f"Error: {e}")
        return render(request, 'main/error.html', {"error_message": str(e)})


@login_required
def compatibility_report_view(request):
    employer = request.user
    search_query = request.GET.get('search', '').strip().lower()

    try:
        # Generate detailed compatibility report
        detailed_report = generate_detailed_compatibility_report(employer=employer)

        if not detailed_report:
            return render(request, 'main/compatibility_report.html', {
                "error_message": "No compatibility data available.",
            })

        # Filter out specific candidates
        filtered_report = [
            entry for entry in detailed_report if entry["Candidate"] != "vinaybharadwaj"
        ]

        # Extract all unique companies and employees
        all_companies = list({entry["Job"].split(" - ")[0] for entry in filtered_report})
        all_employees = [{'full_name': e.full_name, 'user_id': e.user_id} for e in UserProfile.objects.filter(role='employee')]

        # Create the full similarity matrix
        full_similarity_df = pd.DataFrame(index=all_companies, columns=[e['full_name'] for e in all_employees], dtype=object).fillna({"score": 0, "job_id": None})

        for entry in filtered_report:
            company = entry["Job"].split(" - ")[0]
            employee_full_name = entry["Candidate"]
            score = entry["Overall Compatibility"]
            job_id = entry.get("job_id")
            user_id = next((e["user_id"] for e in all_employees if e["full_name"] == employee_full_name), None)
            full_similarity_df.at[company, employee_full_name] = {"score": score, "job_id": job_id, "user_id": user_id}

        # Handle search query filtering
        filtered_chart_df = full_similarity_df
        filtered_similarity_matrix = None
        filtered_message = None

        if search_query:
            # Filter companies based on the search query
            filtered_companies = [c for c in all_companies if search_query in c.lower()]
            if filtered_companies:
                filtered_chart_df = full_similarity_df.loc[filtered_companies]
                filtered_similarity_matrix = filtered_chart_df.to_dict(orient="index")
            else:
                filtered_message = f"No company found matching '{search_query}'. Displaying all companies."
                filtered_chart_df = full_similarity_df

        # Generate the bar chart
        bar_chart_path = None
        if not filtered_chart_df.empty:
            chart_path = os.path.join(settings.STATICFILES_DIRS[0], f"reports/clustered_{employer.userprofile.full_name}.png")
            generate_clustered_bar_chart(filtered_chart_df.applymap(lambda x: x["score"] if isinstance(x, dict) else 0).T, chart_path)
            bar_chart_path = f"/static/reports/clustered_{employer.userprofile.full_name}.png"

        # Simplify matrices for rendering
        def simplify_matrix(matrix):
            return {company: {employee: (data if isinstance(data, dict) else {"score": 0, "job_id": None, "user_id": next((e["user_id"] for e in all_employees if e["full_name"] == employee), None)}) for employee, data in employees.items()} for company, employees in matrix.items()}

        full_similarity_matrix = simplify_matrix(full_similarity_df.to_dict(orient="index"))
        if filtered_similarity_matrix is not None:
            filtered_similarity_matrix = simplify_matrix(filtered_similarity_matrix)

        # Render the template
        return render(request, 'main/compatibility_report.html', {
            "bar_chart_path": bar_chart_path,
            "full_similarity_matrix": full_similarity_matrix,
            "filtered_similarity_matrix": filtered_similarity_matrix,
            "search_query": search_query,
            "filtered_message": filtered_message,
            "employees": all_employees,
        })

    except Exception as e:
        logger.error(f"Error: {e}")
        return render(request, 'main/error.html', {"error_message": "An error occurred while processing the request."})




@login_required
@employee_required
def employee_report_view(request):
    try:
        current_employee = request.user.userprofile.full_name.strip()  # Use full_name instead of username

        # Generate detailed report
        detailed_report = generate_employee_compatibility_report(request.user)
        if not detailed_report:
            raise ValueError("No compatibility data found for this employee.")

        # Build similarity DataFrame
        similarity_df = pd.DataFrame(
            {entry["Job"]: entry["Overall Compatibility"] for entry in detailed_report},
            index=[entry["Candidate"] for entry in detailed_report],
        ).T

        if current_employee not in similarity_df.index:
            raise ValueError(f"No compatibility data available for employee: {current_employee}")

        # Get recommendations
        recommendations = recommend_top_jobs(similarity_df, Job.objects.all())
        employee_recommendations = recommendations.get(current_employee, [])

        # Generate clustered chart
        clustered_chart_path = os.path.join(REPORTS_DIR, f'{current_employee.replace(" ", "_")}_compatibility_clustered_chart.png')
        if not os.path.exists(REPORTS_DIR):
            os.makedirs(REPORTS_DIR)
        generate_employee_clustered_chart(similarity_df.loc[current_employee], clustered_chart_path)

        return render(request, 'main/employee_report.html', {
            'recommendations': employee_recommendations,
            'clustered_chart_path': f'/static/reports/{os.path.basename(clustered_chart_path)}',
            'employee_name': current_employee,
        })

    except Exception as e:
        logger.error(f"Error in employee_report_view: {e}")
        return render(request, 'main/error.html', {"error_message": str(e)})


from django.utils.text import slugify

# Generate URL-safe job name
def get_url_safe_job_name(company_name):
    return company_name.replace(' ', '-').replace(',', 'comma').replace('.', 'period')


@login_required
@employee_required
def view_employee_compatibility_report(request, job_name):
    """
    Fetch and display a detailed compatibility report for a specific job.
    """
    try:
        current_employee = request.user.userprofile.full_name.strip()  # Use full_name instead of username

        # Replace hyphens with spaces and handle special characters to match the original company names
        job_name = job_name.replace('-', ' ').replace('comma', ',').replace('period', '.')

        # Generate the detailed report for the employee
        detailed_report = generate_employee_compatibility_report(request.user)

        # Log available jobs for debugging
        logger.debug(f"Available jobs: {[entry['Job'] for entry in detailed_report]}")
        logger.debug(f"Requested Job: {job_name}")

        # Find the job report by comparing job names in a case-insensitive manner
        job_report = next((entry for entry in detailed_report if job_name.lower() == entry["Job"].lower()), None)

        if not job_report:
            logger.error(f"No compatibility data found for the job: {job_name}")
            raise ValueError(f"No compatibility data found for the job: {job_name}")

        # Extract company name
        company_name = job_report["Job"].strip()

        # Log the company for debugging
        logger.debug(f"Looking for job details for company: {company_name}")

        # Fetch the job details from the database using a case-insensitive match
        job_details = Job.objects.filter(
            company_name__iexact=company_name  # Exact match for company name only
        ).values(
            'education', 'skills', 'experience'
        ).first()

        if not job_details:
            logger.error(f"No job details found for company: {company_name}")
            raise ValueError(f"No job details found for company: {company_name}")

        # Prepare job details dictionary
        job_details_dict = {
            "education": job_details.get("education", "Not Specified"),
            "skills": [skill.strip() for skill in (job_details.get("skills", "") or "").split(",")],
            "experience": job_details.get("experience", "0 - 0")
        }

        # Extract details for rendering
        compatibility_matrix = job_report.get("Criteria", [])
        overall_compatibility = {
            "education_compatibility": job_report.get("Education Compatibility", 0),
            "skills_compatibility": job_report.get("Skills Compatibility", 0),
            "experience_compatibility": job_report.get("Experience Compatibility", 0),
            "overall_compatibility": job_report.get("Overall Compatibility", 0),
        }

        # Generate recommendations using the existing function
        recommendations = generate_recommendations(compatibility_matrix, job_details=job_details_dict)

        # Render the compatibility report
        return render(request, 'main/employee_side_compatibility_display.html', {
            "employee_name": current_employee,
            "company_name": company_name,
            "compatibility_matrix": compatibility_matrix,
            "overall_compatibility": overall_compatibility,
            "recommendations": recommendations,
        })

    except Exception as e:
        logger.error(f"Error in view_employee_compatibility_report: {e}")
        return render(request, 'main/error.html', {"error_message": str(e)})

from django.utils.text import slugify

@login_required
@employee_required
def view_compatibility_scores(request):
    try:
        current_employee = request.user.userprofile.full_name.strip()  # Use full_name instead of username

        # Generate detailed report
        detailed_report = generate_employee_compatibility_report(request.user)
        if not detailed_report:
            raise ValueError("No compatibility data found for this employee.")

        # Build similarity DataFrame
        similarity_df = pd.DataFrame(
            {entry["Job"]: entry["Overall Compatibility"] / 100 for entry in detailed_report},
            index=[current_employee]
        ).T

        # Sort compatibility scores
        compatibility_scores = similarity_df[current_employee].sort_values(ascending=False).to_dict()

        # Prepare data with slugs and details, only storing the company name and score
        slugified_scores = {}
        for job, score in compatibility_scores.items():
            company_name = job.split(" - ")[0].strip()  # Extract the company name only

            # Generate URL-safe job name
            url_safe_name = company_name.replace(' ', '-').replace(',', 'comma').replace('.', 'period')

            # Ensure only valid job entries are included
            if Job.objects.filter(company_name__iexact=company_name).exists():
                slugified_scores[slugify(company_name)] = {
                    'company_name': company_name,
                    'url_safe_name': url_safe_name,
                    'score': score
                }
            else:
                logger.error(f"No compatibility data found for the job: {company_name}")

        # Store top 10 compatibility scores in the session
        top_10_scores = list(slugified_scores.items())[:10]
        request.session['top_10_compatibility_scores'] = top_10_scores

        return render(request, 'main/view_compatibility_scores.html', {
            'compatibility_scores': slugified_scores,
            'employee_name': current_employee,  # Update the context with full_name
        })

    except Exception as e:
        logger.error(f"Error in view_compatibility_scores: {e}")
        return render(request, 'main/error.html', {"error_message": str(e)})


@login_required
@employee_required
def download_compatibility_scores(request):
    try:
        # Generate compatibility matrix
        detailed_report = generate_employee_compatibility_report(request.user)
        if not detailed_report:
            raise ValueError("No compatibility data available for download.")

        # Build DataFrame with full names
        similarity_df = pd.DataFrame(
            {entry["Job"]: entry["Overall Compatibility"] for entry in detailed_report},
            index=[request.user.userprofile.full_name]  # Use full_name instead of username
        ).T

        # Prepare CSV response
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="compatibility_scores.csv"'
        similarity_df.to_csv(response, index=True)
        return response

    except Exception as e:
        logger.error(f"Error in download_compatibility_scores: {e}")
        return render(request, 'main/error.html', {"error_message": str(e)})



@login_required
def employer_side_openaiCS(request):
    # Set OpenAI API key
    openai.api_key = settings.OPENAI_API_KEY

    # Excluded user full names
    EXCLUDED_USERS = ["vinaybharadwaj", "admin"]

    # Fetch all posted jobs for the employer
    employer_jobs = Job.objects.filter(employer=request.user)
    compatibility_data = {}
    employees_set = set()  # To collect unique employee full names

    for job in employer_jobs:
        compatibility_data[job.company_name] = {}

        # Fetch employees for compatibility scoring, excluding by full_name
        employees = UserProfile.objects.filter(role='employee').exclude(full_name__in=EXCLUDED_USERS)

        for employee in employees:
            # Initialize entry for each employee
            full_name = employee.full_name.strip()
            employees_set.add(full_name)
            compatibility_data[job.company_name][full_name] = None  # Initialize with None

            try:
                # Call OpenAI API
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are an AI compatibility scorer."},
                        {"role": "user", "content": f"Score compatibility between job: {job.job_description} and employee with skills: {employee.skills}. Return a score from 0 to 100."}
                    ]
                )
                score = response['choices'][0]['message']['content'].strip()

                # Store compatibility score
                compatibility_data[job.company_name][full_name] = float(score)

            except ValueError as e:
                compatibility_data[job.company_name][full_name] = f"Error: {str(e)}"
            except Exception as e:
                compatibility_data[job.company_name][full_name] = f"API Error: {str(e)}"

    return render(request, 'main/employer_side_openaiCS.html', {
        'compatibility_data': compatibility_data,
        'employees_headers': sorted(employees_set),  # Pass sorted employee full names
    })



from django.shortcuts import render, get_object_or_404
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from io import BytesIO

@login_required
def employer_side_openaiCR(request, job, employee):
    # Set OpenAI API key
    openai.api_key = settings.OPENAI_API_KEY

    # Fetch job and employee details
    job_obj = get_object_or_404(Job, company_name=job)
    employee_obj = get_object_or_404(UserProfile, full_name=employee)  # Use full_name to fetch

    # Call OpenAI for detailed analysis
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an AI compatibility analysis expert."},
                {"role": "user", "content": f"Provide a detailed compatibility report between job: {job_obj.job_description} "
                                            f"and employee with skills: {employee_obj.skills}. Include insights and recommendations."}
            ]
        )
        detailed_report = response['choices'][0]['message']['content'].strip()
    except openai.error.OpenAIError as e:
        return render(request, 'main/error.html', {"error_message": f"OpenAI API error: {str(e)}"})

    if 'download' in request.GET:
        # Generate PDF
        template = 'main/pdf_template.html'
        context = {
            'job': job,
            'employee_full_name': employee_obj.full_name,  # Ensure this is used for rendering
            'compatibility_details': detailed_report
        }
        html = render_to_string(template, context)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Compatibility_Report_{job}_{employee}.pdf"'

        pisa_status = pisa.CreatePDF(html, dest=response)
        if pisa_status.err:
            return HttpResponse("PDF generation failed")
        return response

    return render(request, 'main/employer_side_openaiCR.html', {
        'job': job,
        'employee_full_name': employee_obj.full_name,  # Correctly pass full name
        'compatibility_details': detailed_report
    })



@login_required
def employee_side_openaiCS(request):
    # Set OpenAI API key
    openai.api_key = settings.OPENAI_API_KEY

    # Get the logged-in user's profile
    try:
        logged_in_employee = UserProfile.objects.get(user=request.user, role='employee')
    except UserProfile.DoesNotExist:
        return render(request, 'main/employee_side_openaiCS.html', {
            'error': 'You are not authorized to view this page.',
        })

    # Fetch all jobs
    jobs = Job.objects.all()
    compatibility_data = {}

    for job in jobs:
        try:
            # Call OpenAI API to generate compatibility score for the logged-in employee
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an AI compatibility scorer."},
                    {"role": "user", "content": f"Score compatibility between job: {job.job_description} "
                                                f"and employee with skills: {logged_in_employee.skills}. "
                                                f"Return a score from 0 to 100."}
                ]
            )
            raw_score = response['choices'][0]['message']['content'].strip()
            # Extract numeric score using regex
            match = re.search(r'(\d+\.?\d*)', raw_score)
            if match:
                score = float(match.group(1))
                compatibility_data[job.company_name] = score
            else:
                compatibility_data[job.company_name] = "Error: No valid score found in the API response."

        except Exception as e:
            logger.error(f"OpenAI API error for job {job.company_name}: {str(e)}")
            compatibility_data[job.company_name] = "Error: Unable to fetch compatibility score."

    # Render the template with compatibility data for the logged-in employee
    return render(request, 'main/employee_side_openaiCS.html', {
        'compatibility_data': compatibility_data,
        'employee': logged_in_employee.full_name,  # Pass the logged-in employee's full_name
    })



from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from io import BytesIO

from urllib.parse import unquote

@login_required
def employee_side_openaiCR(request, job, employee):
    employee = unquote(employee)  # Decode URL-encoded string

    job_obj = get_object_or_404(Job, company_name=job)
    employee_profile = get_object_or_404(UserProfile, full_name=employee)

    # Call OpenAI to generate detailed compatibility analysis
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an AI compatibility scorer providing detailed analysis."},
                {"role": "user", "content": f"Generate a detailed compatibility report between job: {job_obj.job_description} "
                                            f"and employee with skills: {employee_profile.skills}. Highlight strengths, gaps, "
                                            f"and areas for improvement. Provide recommendations."}
            ]
        )
        detailed_report = response['choices'][0]['message']['content'].strip()
    except Exception as e:
        detailed_report = f"Error generating report: {str(e)}"

    # Handle PDF download if requested
    if "download" in request.GET:
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Add content to the PDF
            story.append(Paragraph("<b>Detailed Compatibility Report</b>", styles['Title']))
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"<b>Job:</b> {job_obj.company_name}", styles['Normal']))
            story.append(Paragraph(f"<b>Employee:</b> {employee_profile.full_name}", styles['Normal']))
            story.append(Spacer(1, 12))

            # Add report sections
            sections = ["Detailed Analysis"]
            for section in sections:
                story.append(Paragraph(f"<b>{section}:</b>", styles['Heading2']))
                story.append(Paragraph(detailed_report, styles['Normal']))
                story.append(Spacer(1, 12))

            doc.build(story)
            buffer.seek(0)

            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="Compatibility_Report_{job_obj.company_name}_{employee_profile.full_name}.pdf"'
            return response
        except Exception as e:
            return render(request, 'main/error.html', {"error_message": f"PDF generation failed: {str(e)}"})

    # Render compatibility report to the template
    return render(request, 'main/employee_side_openaiCR.html', {
        'company_name': job_obj.company_name,
        'employee_name': employee_profile.full_name,  # Correctly pass full name
        'detailed_report': detailed_report,
    })

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from .models import EmployeeCertification
from .forms import EmployeeCertificationForm

@login_required
@user_passes_test(lambda u: hasattr(u, 'userprofile'))
def manage_certifications(request):
    """
    View to manage employee certifications
    """
    # Get all certifications for the current user
    certifications = EmployeeCertification.objects.filter(user=request.user)
    
    if request.method == 'POST':
        form = EmployeeCertificationForm(request.POST, request.FILES)
        if form.is_valid():
            certification = form.save(commit=False)
            certification.user = request.user
            certification.save()
            
            messages.success(request, f"Certification '{certification.certificate_name}' added successfully!")
            return redirect('manage_certifications')
    else:
        form = EmployeeCertificationForm()
    
    return render(request, 'main/manage_certifications.html', {
        'certifications': certifications,
        'form': form
    })

@login_required
@user_passes_test(lambda u: hasattr(u, 'userprofile'))
def delete_certification(request, certification_id):
    """
    View to delete a specific certification
    """
    certification = get_object_or_404(
        EmployeeCertification, 
        id=certification_id, 
        user=request.user
    )
    
    if request.method == 'POST':
        certification.delete()
        messages.success(request, f"Certification '{certification.certificate_name}' deleted successfully!")
        return redirect('manage_certifications')
    
    return render(request, 'main/confirm_delete_certification.html', {
        'certification': certification
    })

# Employee Compatibility Report using OpenAI

import tempfile
import logging
import os
import PyPDF2
import pytesseract
import docx2txt
from PIL import Image
from openai import OpenAI, APIError, AuthenticationError, RateLimitError

# Configure logging
logger = logging.getLogger(__name__)

# Tesseract Path Configuration
TESSERACT_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
if os.path.exists(TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
else:
    logger.warning(f"Tesseract not found at {TESSERACT_PATH}. Please check the installation.")

def extract_text_from_file(file):
    """
    Extract text from various file types
    Supports PDF, DOCX, JPG, PNG
    """
    # Determine file extension
    file_ext = os.path.splitext(file.name)[1].lower()
    
    try:
        # Temporary file to work with
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            for chunk in file.chunks():
                temp_file.write(chunk)
            temp_file.close()
        
        # Extract text based on file type
        text = ''
        if file_ext == '.pdf':
            try:
                with open(temp_file.name, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text = ''
                    for page in pdf_reader.pages:
                        text += page.extract_text() or ''
            except Exception as pdf_error:
                logger.error(f"PDF extraction error: {pdf_error}")
        
        elif file_ext == '.docx':
            try:
                text = docx2txt.process(temp_file.name) or ''
            except Exception as docx_error:
                logger.error(f"DOCX extraction error: {docx_error}")
        
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            try:
                # Use Tesseract OCR for image files
                image = Image.open(temp_file.name)
                text = pytesseract.image_to_string(image) or ''
            except Exception as ocr_error:
                logger.error(f"OCR extraction error: {ocr_error}")
        
        else:
            text = ''
        
        # Clean up temporary file
        os.unlink(temp_file.name)
        
        return text.strip()
    
    except Exception as e:
        logger.error(f"Error extracting text: {e}")
        return ''

# Updated employee_compatibility view function
@login_required
def employee_compatibility(request):
    context = {}
    
    if request.method == 'POST':
        # Validate OpenAI API key
        openai_api_key = getattr(settings, 'OPENAI_API_KEY', None)
        if not openai_api_key:
            messages.error(request, "OpenAI API key is not configured.")
            return render(request, 'main/employee_compatibility.html', context)
        
        # Set up OpenAI client
        client = OpenAI(api_key=settings.OPENAI_API_KEY)

        # Validate files - using the multi-file upload method
        resumes = request.FILES.getlist('resumes')
        job_description_file = request.FILES.get('job_description')
        
        # Check if files are uploaded
        if not resumes:
            messages.error(request, "Please upload at least one resume.")
            return render(request, 'main/employee_compatibility.html', context)
            
        if not job_description_file:
            messages.error(request, "Please upload a job description.")
            return render(request, 'main/employee_compatibility.html', context)
        
        # Log file information
        logger.info(f"Received {len(resumes)} resume(s) through multi-file upload:")
        for i, resume in enumerate(resumes):
            logger.info(f"  Resume {i+1}: {resume.name} ({resume.size} bytes)")
        logger.info(f"Job description: {job_description_file.name} ({job_description_file.size} bytes)")
        
        # Validate file sizes (5MB limit)
        MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
        oversized_files = []
        
        for resume in resumes:
            if resume.size > MAX_FILE_SIZE:
                oversized_files.append(resume.name)
                
        if job_description_file.size > MAX_FILE_SIZE:
            oversized_files.append(job_description_file.name)
            
        if oversized_files:
            message = f"The following files exceed the 5MB limit: {', '.join(oversized_files)}"
            messages.error(request, message)
            return render(request, 'main/employee_compatibility.html', context)
        
        # Extract text from files
        job_description_text = extract_text_from_file(job_description_file)
        resume_texts = []
        resume_names = []
        
        for resume in resumes:
            text = extract_text_from_file(resume)
            resume_texts.append(text)
            resume_names.append(resume.name)
            
        # Generate compatibility report using OpenAI
        try:
            compatibility_reports = []
            for i, (resume_text, resume_name) in enumerate(zip(resume_texts, resume_names), 1):
                # Skip empty resumes
                if not resume_text.strip():
                    compatibility_reports.append(f"<div class='alert alert-warning report-section mb-5'><h3 class='mt-4 mb-3'>Analysis for Resume {i}: {resume_name}</h3><p>No extractable text found in this file.</p></div>")
                    continue
                
                # Add resume name to the report - use semantic HTML5 for better PDF rendering
                compatibility_reports.append(f"<div class='report-section mb-5'><h3 class='mt-4 mb-3'>Analysis for Resume {i}: {resume_name}</h3>")
                
                # Updated prompt to include Skills and Experience Compatibility Scores
                prompt = f"""Perform a detailed compatibility analysis between the following job description and resume:

Job Description:
{job_description_text}

Resume {i} ({resume_name}):
{resume_text}

Provide a comprehensive analysis with the following details:
1. Overall Compatibility Score (0-100%)
2. Skills Compatibility Score (0-100%)
3. Experience Compatibility Score (0-100%)
4. Skills Match Analysis - List skills match points as bullet points (at least 3-4 bullets)
5. Experience Relevance - List experience relevance as bullet points (at least 3-4 bullets)
6. Key Strengths - List at least 3 key strengths
7. Potential Gaps - List any potential gaps
8. Recommendation - List recommendations as bullet points (at least 3-4 bullets)

Format the response using this exact HTML structure:
<table class="table table-bordered table-striped">
  <tbody>
    <tr>
      <th width="25%">Overall Compatibility Score</th>
      <td><strong>XX%</strong></td>
    </tr>
    <tr>
      <th width="25%">Skills Compatibility Score</th>
      <td><strong>XX%</strong></td>
    </tr>
    <tr>
      <th width="25%">Experience Compatibility Score</th>
      <td><strong>XX%</strong></td>
    </tr>
    <tr>
      <th>Skills Match Analysis</th>
      <td>
        <ul>
          <li>[Skills match point 1]</li>
          <li>[Skills match point 2]</li>
          <li>[Skills match point 3]</li>
          <li>[Skills match point 4 if applicable]</li>
        </ul>
      </td>
    </tr>
    <tr>
      <th>Experience Relevance</th>
      <td>
        <ul>
          <li>[Experience relevance point 1]</li>
          <li>[Experience relevance point 2]</li>
          <li>[Experience relevance point 3]</li>
          <li>[Experience relevance point 4 if applicable]</li>
        </ul>
      </td>
    </tr>
    <tr>
      <th>Key Strengths</th>
      <td>
        <ul>
          <li>[Strength 1]</li>
          <li>[Strength 2]</li>
          <li>[Strength 3]</li>
        </ul>
      </td>
    </tr>
    <tr>
      <th>Potential Gaps</th>
      <td>
        <ul>
          <li>[Gap 1]</li>
          <li>[Gap 2]</li>
        </ul>
      </td>
    </tr>
    <tr>
      <th>Recommendation</th>
      <td>
        <ul>
          <li>[Recommendation point 1]</li>
          <li>[Recommendation point 2]</li>
          <li>[Recommendation point 3]</li>
          <li>[Recommendation point 4 if applicable]</li>
        </ul>
      </td>
    </tr>
  </tbody>
</table>

Do not modify this table structure. Keep all <th> and <td> elements exactly as shown."""
                
                try:
                    # Updated API call for newer OpenAI client
                    response = client.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": "You are a professional HR analyst doing a resume-job description compatibility check. Format your response using the exact HTML structure provided by the user. Do not deviate from the requested format."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=1500
                    )
                    
                    # Add the response to our reports collection
                    compatibility_reports.append(response.choices[0].message.content)
                    # Close the section
                    compatibility_reports.append("</div>")
                
                except APIError as api_error:
                    logger.error(f"OpenAI API Error for Resume {i} ({resume_name}): {str(api_error)}")
                    compatibility_reports.append(f"<div class='alert alert-danger'><p>API Error for Resume {i} ({resume_name}): {str(api_error)}</p></div>")
                
                except RateLimitError as rate_error:
                    logger.error(f"Rate limit exceeded for Resume {i} ({resume_name}): {str(rate_error)}")
                    compatibility_reports.append(f"<div class='alert alert-warning'><p>Rate limit exceeded. Please try again in a moment for Resume {i} ({resume_name}).</p></div>")
                
                except AuthenticationError as auth_error:
                    logger.error(f"Authentication error: {str(auth_error)}")
                    compatibility_reports.append(f"<div class='alert alert-danger'><p>Authentication error. Please check your API key configuration.</p></div>")
                
                except Exception as e:
                    logger.error(f"Error for Resume {i} ({resume_name}): {str(e)}")
                    compatibility_reports.append(f"<div class='alert alert-danger'><p>Error generating report for Resume {i} ({resume_name}): {str(e)}</p></div>")
            
            # Combine reports
            context['compatibility_report'] = "".join(compatibility_reports)
            
            # Make sure Django doesn't escape the HTML
            from django.utils.safestring import mark_safe
            context['compatibility_report'] = mark_safe(context['compatibility_report'])
            
            context['report_count'] = len(resumes)
        
        except Exception as e:
            logger.error(f"Error generating compatibility report: {str(e)}")
            messages.error(request, f"Error generating compatibility report: {str(e)}")
    
    return render(request, 'main/employee_compatibility.html', context)
